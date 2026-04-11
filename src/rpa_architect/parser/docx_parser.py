"""DOCX parser implementation using python-docx.

Extracts text sections with heading hierarchy, tables, and embedded images
from Word-format Process Design Documents. Preserves formatting context
(bold, italic) as semantic hints for LLM extraction.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from rpa_architect.parser.base import PddContent, PddSection, PddTable

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse DOCX-format PDD documents using python-docx.

    Extracts text with heading hierarchy, tables with headers, and
    embedded images. Bold and italic text is annotated with markers
    to preserve formatting context for downstream LLM processing.
    """

    # Mapping from python-docx heading style level to our level
    HEADING_STYLE_PREFIX = "Heading"
    BOLD_MARKER_START = "**"
    BOLD_MARKER_END = "**"
    ITALIC_MARKER_START = "_"
    ITALIC_MARKER_END = "_"

    def parse(self, path: Path) -> PddContent:
        """Parse a DOCX file and extract structured content.

        Args:
            path: Path to the DOCX file.

        Returns:
            PddContent with sections, tables, images, and metadata.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid DOCX.
        """
        try:
            import docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            ) from exc

        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        if path.suffix.lower() not in (".docx", ".doc"):
            raise ValueError(f"Expected a .docx file, got: {path.suffix}")

        doc = docx.Document(str(path))

        sections: list[PddSection] = []
        tables: list[PddTable] = []
        images: list[bytes] = []
        metadata: dict[str, Any] = {}

        # Extract document metadata from core properties
        metadata.update(self._extract_metadata(doc))

        # Process document body: paragraphs and tables in order
        sections = self._extract_sections(doc)
        tables = self._extract_tables(doc)
        images = self._extract_images(doc)

        return PddContent(
            sections=sections,
            tables=tables,
            images=images,
            metadata=metadata,
        )

    def _extract_metadata(self, doc: Any) -> dict[str, Any]:
        """Extract document metadata from core properties."""
        metadata: dict[str, Any] = {}

        try:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created"] = str(props.created)
            if props.modified:
                metadata["modified"] = str(props.modified)
            if props.subject:
                metadata["subject"] = props.subject
            if props.category:
                metadata["category"] = props.category
            if props.keywords:
                metadata["keywords"] = props.keywords
            if props.comments:
                metadata["comments"] = props.comments
            if props.revision:
                metadata["revision"] = str(props.revision)
        except Exception:
            logger.debug("Failed to extract DOCX core properties")

        return metadata

    def _extract_sections(self, doc: Any) -> list[PddSection]:
        """Extract text sections from the document preserving heading hierarchy.

        Paragraphs are grouped by their heading hierarchy. Formatting (bold,
        italic) is preserved using markdown-style markers.
        """
        sections: list[PddSection] = []
        current_heading: str = ""
        current_level: int = 1
        current_content_lines: list[str] = []

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            text = self._extract_paragraph_text(paragraph)

            if not text.strip():
                # Preserve paragraph breaks in content
                if current_content_lines and current_content_lines[-1] != "":
                    current_content_lines.append("")
                continue

            # Check if this paragraph is a heading
            heading_level = self._get_heading_level(style_name)

            if heading_level is not None:
                # Save previous section
                if current_heading or current_content_lines:
                    content = "\n".join(current_content_lines).strip()
                    sections.append(
                        PddSection(
                            title=current_heading,
                            content=content,
                            level=current_level,
                            page_number=0,  # DOCX doesn't provide page numbers directly
                        )
                    )

                current_heading = text.strip()
                current_level = heading_level
                current_content_lines = []
            else:
                current_content_lines.append(text)

        # Save the last section
        if current_heading or current_content_lines:
            content = "\n".join(current_content_lines).strip()
            sections.append(
                PddSection(
                    title=current_heading,
                    content=content,
                    level=current_level,
                    page_number=0,
                )
            )

        return sections

    def _extract_paragraph_text(self, paragraph: Any) -> str:
        """Extract text from a paragraph, preserving bold/italic formatting.

        Bold text is wrapped in **, italic in _, to preserve semantic
        emphasis for LLM processing.
        """
        parts: list[str] = []

        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            if run.bold and run.italic:
                parts.append(
                    f"{self.BOLD_MARKER_START}{self.ITALIC_MARKER_START}"
                    f"{text}"
                    f"{self.ITALIC_MARKER_END}{self.BOLD_MARKER_END}"
                )
            elif run.bold:
                parts.append(f"{self.BOLD_MARKER_START}{text}{self.BOLD_MARKER_END}")
            elif run.italic:
                parts.append(f"{self.ITALIC_MARKER_START}{text}{self.ITALIC_MARKER_END}")
            else:
                parts.append(text)

        return "".join(parts)

    def _get_heading_level(self, style_name: str) -> int | None:
        """Determine heading level from a paragraph style name.

        Returns:
            Heading level (1-6) or None if it's not a heading.
        """
        if not style_name:
            return None

        # Standard heading styles: "Heading 1", "Heading 2", etc.
        if style_name.startswith(self.HEADING_STYLE_PREFIX):
            suffix = style_name[len(self.HEADING_STYLE_PREFIX) :].strip()
            try:
                level = int(suffix)
                return min(max(level, 1), 6)
            except ValueError:
                pass

        # Title style = level 1
        if style_name.lower() in ("title",):
            return 1

        # Subtitle = level 2
        if style_name.lower() in ("subtitle",):
            return 2

        return None

    def _extract_tables(self, doc: Any) -> list[PddTable]:
        """Extract all tables from the document."""
        extracted_tables: list[PddTable] = []

        for table in doc.tables:
            rows_data: list[list[str]] = []

            for row in table.rows:
                cell_texts: list[str] = []
                for cell in row.cells:
                    # Extract text from all paragraphs in the cell
                    cell_text = "\n".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    cell_texts.append(cell_text)
                rows_data.append(cell_texts)

            if not rows_data:
                continue

            # First row as headers
            headers = rows_data[0]
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            # Try to detect caption from the paragraph before the table
            caption = self._find_table_caption(doc, table)

            extracted_tables.append(
                PddTable(
                    headers=headers,
                    rows=data_rows,
                    caption=caption,
                    page_number=0,
                )
            )

        return extracted_tables

    def _find_table_caption(self, doc: Any, table: Any) -> str:
        """Try to find a caption for a table by looking at surrounding paragraphs.

        Looks for paragraphs with 'Caption' style or 'Table' prefix near the table.
        """
        try:
            # Walk through the document body elements to find the table
            from docx.oxml.ns import qn

            table_element = table._tbl
            body = doc.element.body

            for i, child in enumerate(body):
                if child is table_element and i > 0:
                    # Check previous element
                    prev = body[i - 1]
                    if prev.tag == qn("w:p"):
                        _text = prev.text or ""
                        # Clean up the text
                        full_text = "".join(
                            node.text or ""
                            for node in prev.iter()
                            if node.text
                        )
                        if full_text.strip():
                            return full_text.strip()
                    break
        except Exception:
            pass

        return ""

    def _extract_images(self, doc: Any) -> list[bytes]:
        """Extract all embedded images from the document."""
        extracted_images: list[bytes] = []

        try:

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        if image_bytes and len(image_bytes) > 100:
                            extracted_images.append(image_bytes)
                    except Exception:
                        logger.debug("Failed to extract an image from DOCX")
                        continue
        except Exception:
            logger.debug("Failed to process images in DOCX")

        return extracted_images
